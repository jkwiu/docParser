from os import walk
import re
import os
from docx.shared import Inches
from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


def make_rows_bold(*rows):
    for row in rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True


def color_row(row=0):
    'make row of cells background colored, defaults to column header row'
    row = table.rows[row]
    for cell in row.cells:
        shading_elm_2 = parse_xml(
            r'<w:shd {} w:fill="1F5C8B"/>'.format(nsdecls('w')))
        cell._tc.get_or_add_tcPr().append(shading_elm_2)


path = './code/'

_, _, fileNames = next(walk(path))

for fileName in fileNames:
    fileName = fileName.replace('.cs', '')
    if os.path.exists('./document/'+fileName+'.docx'):
        os.remove('./document/'+fileName+'.docx')
    else:
        print('The file does not exist')

    with open('./code/'+fileName+'.cs', encoding='utf8') as f:
        lines = f.read().split("\n")

    # 초기 세팅
    initialStartLine = 0
    dic = {}
    params = {}
    smryKeyword = ''
    smryReadOn = False
    paramSmryReadOn = False

    # initialStartLine 찾기
    className = ''
    for i in range(0, len(lines)):
        line = lines[i].lstrip()
        classLine = line.split(' ')
        if classLine[0] == 'public' and classLine[1] == 'class':
            className = classLine[2]
            break
    for i in range(0, len(lines)):
        consFuncName = ''
        line = lines[i].lstrip()
        consLine = line.split(' ')
        if len(consLine) > 1:
            for c in consLine[1]:
                consFuncName += c
                if consFuncName == className:
                    initialStartLine = i + 1
                    break
    for i in range(initialStartLine, len(lines)):
        # Summary Keyword 추출
        line = lines[i].lstrip().split(' ')
        if smryReadOn:
            if lines[i] != '':
                smryKeyword = lines[i]
                smryKeyword = smryKeyword.replace(
                    '///', '').replace('</summary>', '').lstrip()
                smryReadOn = False
        # Param Summary 추출
        if len(line) > 1:
            if line[1] == '<summary>':
                smryReadOn = True
            elif line[1] == '<param':
                paramNameExp = '\"[a-zA-Z]+\"'
                paramNameList = re.findall(paramNameExp, lines[i])
                if len(paramNameList) > 0:
                    paramName = paramNameList[0].replace("\"", "")
                    paramSmry = lines[i].lstrip().replace(
                        '<param name="' + paramName+'">', '').replace('</param>', '').replace('///', '')
                    params[paramName] = {
                        'Param Summary': paramSmry,
                    }

        # Function 추출
        if line[0] == 'public' and line[1] != 'class':
            funcName = ''
            returnName = ''
            paramType = ''
            functionNameIdx = 0
            for idx in range(0, len(line)):
                functionLines = line[idx]
                for key in functionLines:
                    if key == "(":
                        functionNameIdx = idx
                        break
            # Function Name 추출
            funcNames = line[functionNameIdx]
            for key in funcNames:
                if key == "(":
                    break
                funcName += key
            # Return Name 추출
            for idx in range(1, functionNameIdx):
                returnName += line[idx] + ' '
            # Param Type 추출
            for idx in range(functionNameIdx, len(line)):
                paramType += line[idx] + ' '
            paramType = paramType.replace(funcName + '(', '').replace(')', '')
            paramsList = paramType.split(',')
            for idx in range(0, len(paramsList)):
                paramsList[idx] = paramsList[idx].lstrip().rstrip()
            for key in paramsList:
                paramPair = key.split(' ')
                if len(paramPair) > 1:
                    if paramPair[1] in params:
                        params[paramPair[1]]['Type'] = paramPair[0]

            # 전체 키워드 취합
            dic[funcName] = {
                'Summary': smryKeyword.lstrip().rstrip(),
                'Returns': returnName.lstrip().rstrip(),
                'Params': params,
            }
            params = {}

            # word 문서 작성
            doc = Document()

            # 테이블 그리기
            for name in dic:
                parmLength = 0
                for paramName in dic[name]['Params']:
                    if 'Type' in dic[name]['Params'][paramName]:
                        parmLength += 1
                # rowLength = len(dic[name]['Params'])+4
                table = doc.add_table(
                    rows=parmLength+4, cols=3, style='TableGrid')
                # 기본 테이블 설정
                table.cell(0, 1).merge(table.cell(0, 2))
                table.cell(1, 1).merge(table.cell(1, 2))
                table.cell(2, 1).merge(table.cell(2, 2))
                cell_1 = table.rows[0].cells
                cell_2 = table.rows[1].cells
                cell_3 = table.rows[2].cells
                cell_4 = table.rows[3].cells

                cell_1[0].text = "Name"
                cell_2[0].text = "Summary"
                cell_3[0].text = "Returns"
                cell_4[0].text = "Param Name"
                cell_4[1].text = "Type"
                cell_4[2].text = "Param Summary"
                make_rows_bold(table.rows[0], table.rows[1],
                               table.rows[2], table.rows[3])

                # 값 입력
                cell_1[1].text = name
                cell_2[1].text = dic[name]['Summary']
                cell_3[1].text = dic[name]['Returns']
                i = 0
                for paramName in dic[name]['Params']:
                    if 'Type' in dic[name]['Params'][paramName]:
                        cell = table.rows[i+4].cells
                        cell[0].text = paramName
                        cell[1].text = dic[name]['Params'][paramName]['Type']
                        cell[2].text = dic[name]['Params'][paramName]['Param Summary']
                        i += 1
                doc.add_paragraph(' ')

            doc.save('./document/'+fileName+'.docx')
